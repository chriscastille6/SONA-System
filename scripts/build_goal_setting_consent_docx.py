#!/usr/bin/env python3
"""Rebuild goal-setting consent .docx files from .txt sources with Nicholls branding."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Inches, Pt

REPO_ROOT = Path(__file__).resolve().parent.parent
MATERIALS = REPO_ROOT / "apps/studies/assets/irb/goal-setting/materials"
VARIANTS = {
    "pilot": {
        "txt": MATERIALS / "consent_form_pilot.txt",
        "docx": MATERIALS / "consent_form_pilot.docx",
    },
    "main": {
        "txt": MATERIALS / "consent_form.txt",
        "docx": MATERIALS / "consent_form.docx",
    },
}
LOGO_CANDIDATES = [
    MATERIALS / "nicholls_state_university_logo.png",
    Path(
        "/Users/ccastille/.cursor/projects/Users-ccastille-Documents-GitHub-SONA-System/"
        "assets/image-3692823e-13cc-4d71-87c9-875833604dab.png"
    ),
    REPO_ROOT / "static/images/lab_emblem.png",
]

HEADINGS = {
    "Information",
    "Risks",
    "Benefits",
    "Confidentiality of Your Information",
    "Remuneration",
    "Contact",
    "Participation",
    "Feedback and Publication",
    "CONSENT",
}

TITLES = {
    "INFORMED CONSENT STATEMENT/INFORMATION LETTER",
    "A Study in Decision Making",
    "A Study in Decision Making — Pilot Session",
    "A Study in Decision Making — Main Study",
}


def resolve_logo() -> Path:
    for path in LOGO_CANDIDATES:
        if path.is_file():
            return path
    raise FileNotFoundError("No Nicholls logo found in expected locations")


def add_paragraph(doc: Document, text: str, *, heading: bool = False, title: bool = False) -> None:
    para = doc.add_paragraph()
    fmt = para.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = 1.5
    fmt.space_after = Pt(14)
    fmt.space_before = Pt(6 if heading else 0)

    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12 if not title else 14)

    if title:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.bold = True
        fmt.space_after = Pt(10)
    elif heading:
        run.bold = True
        fmt.space_before = Pt(12)
        fmt.space_after = Pt(8)
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def build_docx(variant: str) -> Path:
    paths = VARIANTS[variant]
    txt_path = paths["txt"]
    docx_path = paths["docx"]
    if not txt_path.is_file():
        raise FileNotFoundError(f"Consent source not found: {txt_path}")

    logo_path = resolve_logo()
    content = txt_path.read_text(encoding="utf-8")
    blocks = [b.strip() for b in content.split("\n\n") if b.strip()]

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)

    logo_para = doc.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_para.paragraph_format.space_after = Pt(12)
    logo_run = logo_para.add_run()
    logo_run.add_picture(str(logo_path), width=Inches(2.0))

    for block in blocks:
        if block.startswith("_______"):
            lines = block.splitlines()
            if len(lines) >= 2:
                table = doc.add_table(rows=2, cols=2)
                table.autofit = True
                table.cell(0, 0).text = lines[0].split("\t")[0].strip()
                table.cell(0, 1).text = lines[0].split("\t")[-1].strip() if "\t" in lines[0] else ""
                labels = lines[1].split("\t")
                table.cell(1, 0).text = labels[0].strip() if labels else ""
                table.cell(1, 1).text = labels[-1].strip() if len(labels) > 1 else ""
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.name = "Times New Roman"
                                run.font.size = Pt(12)
            continue

        lines = block.splitlines()
        first = lines[0].strip()
        rest = "\n".join(lines[1:]).strip()

        if first in TITLES and not rest:
            add_paragraph(doc, first, title=True)
            continue

        if first in HEADINGS:
            add_paragraph(doc, first, heading=True)
            if rest:
                add_paragraph(doc, rest)
            continue

        add_paragraph(doc, block)

    doc.save(docx_path)
    print(f"Wrote {docx_path}")
    print(f"Logo: {logo_path}")
    return docx_path


def main() -> None:
    parser = argparse.ArgumentParser(description="Build goal-setting consent .docx from .txt")
    parser.add_argument(
        "--variant",
        choices=sorted(VARIANTS),
        help="Consent variant to build (default: both pilot and main)",
    )
    args = parser.parse_args()
    variants = [args.variant] if args.variant else sorted(VARIANTS)
    for variant in variants:
        build_docx(variant)


if __name__ == "__main__":
    main()

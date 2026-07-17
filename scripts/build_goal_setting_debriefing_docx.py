#!/usr/bin/env python3
"""Rebuild goal-setting debriefing_protocol.docx from debriefing_protocol.txt with Nicholls branding."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Inches, Pt

REPO_ROOT = Path(__file__).resolve().parent.parent
MATERIALS = REPO_ROOT / "apps/studies/assets/irb/goal-setting/materials"
TXT_PATH = MATERIALS / "debriefing_protocol.txt"
DOCX_PATH = MATERIALS / "debriefing_protocol.docx"
LOGO_CANDIDATES = [
    MATERIALS / "nicholls_state_university_logo.png",
    Path(
        "/Users/ccastille/.cursor/projects/Users-ccastille-Documents-GitHub-SONA-System/"
        "assets/image-3692823e-13cc-4d71-87c9-875833604dab.png"
    ),
    REPO_ROOT / "static/images/lab_emblem.png",
]

HEADINGS = {
    "Thank You for Your Participation",
    "Study Purpose and Disclosure",
    "How Your Responses Were Used",
    "Conditions in This Study",
    "What We Are Investigating",
    "Questions or Concerns",
    "Summary of Findings",
}

TITLES = {
    "DEBRIEFING PROTOCOL / APPRECIATION LETTER",
    "A Study in Decision Making",
}

CENTER_LINES = {
    "Nicholls State University",
    "Al Danos College of Business Administration",
}


def resolve_logo() -> Path:
    for path in LOGO_CANDIDATES:
        if path.is_file():
            return path
    raise FileNotFoundError("No Nicholls logo found in expected locations")


def add_paragraph(doc: Document, text: str, *, heading: bool = False, title: bool = False, center: bool = False) -> None:
    para = doc.add_paragraph()
    fmt = para.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = 1.5
    fmt.space_after = Pt(14)
    fmt.space_before = Pt(6 if heading else 0)

    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12 if not title else 14)

    if title or center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if title:
            run.bold = True
            fmt.space_after = Pt(10)
    elif heading:
        run.bold = True
        fmt.space_before = Pt(12)
        fmt.space_after = Pt(8)
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def build_docx() -> Path:
    logo_path = resolve_logo()
    content = TXT_PATH.read_text(encoding="utf-8")
    blocks = [b.strip() for b in content.split("\n\n") if b.strip()]

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)

    logo_para = doc.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_para.paragraph_format.space_after = Pt(12)
    logo_run = logo_para.add_run()
    logo_run.add_picture(str(logo_path), width=Inches(2.0))

    for block in blocks:
        lines = block.splitlines()
        first = lines[0].strip()
        rest = "\n".join(lines[1:]).strip()

        if first in TITLES and not rest:
            add_paragraph(doc, first, title=True)
            continue

        if first in CENTER_LINES and not rest:
            add_paragraph(doc, first, center=True)
            continue

        if first in HEADINGS:
            add_paragraph(doc, first, heading=True)
            if rest:
                add_paragraph(doc, rest)
            continue

        if first.startswith("Dr. ") and "@nicholls.edu" in first and not rest:
            add_paragraph(doc, first, center=True)
            continue

        add_paragraph(doc, block)

    doc.save(DOCX_PATH)
    print(f"Wrote {DOCX_PATH}")
    print(f"Logo: {logo_path}")
    return DOCX_PATH


if __name__ == "__main__":
    build_docx()
